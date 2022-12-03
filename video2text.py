import os
from docx import Document
import requests
from moviepy.editor import AudioFileClip
from speechkit import Session, ShortAudioRecognition
import m3u8
import config 

# here moviepy comes up handy, althouh it looks like it also could be done using just ffmpeg
def video2audio(input_filename='tmp.m4v',output_filename='tmp.wav', samplerate=16000):
    audioclip = AudioFileClip(input_filename)
    audioclip.write_audiofile(output_filename, bitrate='50k', ffmpeg_params=[
        "-ac", "1", "-ar", str(samplerate)])
    # check the size MB
    size_of_the_file = os.path.getsize(output_filename)/1024/1024
    if not size_of_the_file <= 1. :
        print ('File size bigger then 1 MB' )
    # os.remove(input_filename)
    

# downloads bit of the video from the url  : *.ts
def collect_stream_bits(url,filename='tmp.m4v'):
    local = requests.get(url, stream=True)
    if (local.status_code == 200):
        with open(filename, 'wb') as f:
            for chunk in local.iter_content(chunk_size=1024):
                f.write(chunk)
    else:
        print('Error in requests {}'.format(local.status_code))
        
# just shorthand for audio reading
def loadaudio(filename=''):
    with open('/Users/lusigeondzian/Desktop/test/tmp.wav', 'rb') as f:
        data = f.read()
    return data

def write2doc(filename='',data=''):
    document = Document()
    document.add_heading(filename, level=1)
    document.add_paragraph(data)
    document.save('{}.docx'.format(filename))
    


if __name__ == '__main__' :
    #luanch session with yandex cloud
    session = Session.from_yandex_passport_oauth_token(
            config.oauth_token, config.catalog_id)
    # example of url to 'https://somewebsite/streamer/default/storage/api-storage/files/something.mp4/' ;  it has to lead us to chunklist.m3u8 file
    url = 'here goes your url'
    # load infomation about streaming
    m3u8_obj = m3u8.load(url+'chunklist.m3u8')
    # list of all bits *.ts
    playlist = [el['uri'] for el in m3u8_obj.data['segments']]
    # iterate through all bits
    error=[]
    transcribtion=''
    for  item in playlist:
        try:
            url = url+item
            collect_stream_bits(url,filename='tmp.m4v' )
            video2audio(input_filename='tmp.m4v',output_filename='tmp.wav', samplerate=16000)
            data = loadaudio('tmp.wav')
            # the thing is that for  ShortAudioRecongitaion the audio have to be < 1MB
            text = ShortAudioRecognition(session).recognize(
                  data, format='lpcm', sampleRateHertz="16000")
        except:
            error.append(item)
        # there is some definitive problem with punctiation in the audio recogniation
        transcribtion +=  text+'.'+'\n'
    write2doc(filename='trasnscribition', data=transcribtion)
    print('problems occured with {} files from the plyalist' .format(len(error)))